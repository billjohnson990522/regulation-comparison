"""
PDF法规文档解析与转换模块
将PDF格式的法规文档转换为用户友好的Word格式
"""

import os
import re
import json
from typing import List, Dict, Optional, Tuple
from pathlib import Path

# PDF解析库
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# Word生成库
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class RegulationStructure:
    """法规结构化数据"""
    
    def __init__(self):
        self.title: str = ""  # 法规标题
        self.chapters: List[Dict] = []  # 章节列表
    
    def to_dict(self) -> Dict:
        return {
            "title": self.title,
            "chapters": self.chapters
        }
    
    def to_json(self, filepath: str) -> None:
        """导出为JSON格式"""
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(self.to_dict(), f, ensure_ascii=False, indent=2)
        print(f"已导出JSON: {filepath}")


class PDFParser:
    """PDF解析器"""
    
    def __init__(self, use_library: str = "auto"):
        """
        初始化解析器
        
        Args:
            use_library: 使用的解析库，可选 "pymupdf", "pdfplumber", "auto"
        """
        self.library = self._select_library(use_library)
        
    def _select_library(self, use_library: str) -> str:
        """选择PDF解析库"""
        if use_library == "auto":
            if HAS_PYMUPDF:
                return "pymupdf"
            elif HAS_PDFPLUMBER:
                return "pdfplumber"
            else:
                raise ImportError("请安装 PyMuPDF 或 pdfplumber: pip install pymupdf 或 pip install pdfplumber")
        elif use_library == "pymupdf":
            if not HAS_PYMUPDF:
                raise ImportError("请安装 PyMuPDF: pip install pymupdf")
            return "pymupdf"
        elif use_library == "pdfplumber":
            if not HAS_PDFPLUMBER:
                raise ImportError("请安装 pdfplumber: pip install pdfplumber")
            return "pdfplumber"
        else:
            raise ValueError(f"不支持的解析库: {use_library}")
    
    def extract_text(self, pdf_path: str) -> List[Dict]:
        """
        提取PDF文本，按页返回
        
        Returns:
            [{"page": 1, "text": "...", "blocks": [...]}]
        """
        if self.library == "pymupdf":
            return self._extract_with_pymupdf(pdf_path)
        else:
            return self._extract_with_pdfplumber(pdf_path)
    
    def _extract_with_pymupdf(self, pdf_path: str) -> List[Dict]:
        """使用PyMuPDF提取文本"""
        pages = []
        doc = fitz.open(pdf_path)
        
        for page_num, page in enumerate(doc, 1):
            # 获取文本块
            blocks = page.get_text("dict")["blocks"]
            
            text_content = []
            block_list = []
            
            for block in blocks:
                if "lines" in block:
                    block_text = ""
                    for line in block["lines"]:
                        line_text = "".join([span["text"] for span in line["spans"]])
                        block_text += line_text
                    
                    if block_text.strip():
                        text_content.append(block_text)
                        block_list.append({
                            "text": block_text,
                            "bbox": block.get("bbox", [])
                        })
            
            pages.append({
                "page": page_num,
                "text": "\n".join(text_content),
                "blocks": block_list
            })
        
        doc.close()
        return pages
    
    def _extract_with_pdfplumber(self, pdf_path: str) -> List[Dict]:
        """使用pdfplumber提取文本"""
        pages = []
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                
                # 获取文本块
                chars = page.chars
                block_list = []
                
                if chars:
                    # 简单按行分组
                    lines = {}
                    for char in chars:
                        y = round(char["top"], 0)
                        if y not in lines:
                            lines[y] = []
                        lines[y].append(char["text"])
                    
                    for y in sorted(lines.keys()):
                        line_text = "".join(lines[y])
                        if line_text.strip():
                            block_list.append({"text": line_text})
                
                pages.append({
                    "page": page_num,
                    "text": text,
                    "blocks": block_list
                })
        
        return pages


class RegulationParser:
    """法规结构解析器"""
    
    # 标题匹配正则表达式
    CHAPTER_PATTERN = r'^第([一二三四五六七八九十百]+)章[ \s]*(.*)$'  # 第一章 总则
    ARTICLE_PATTERN = r'^第([一二三四五六七八九十百]+)条[ \s]*(.*)$'  # 第一条 为了...
    ITEM_PATTERN = r'^[（(]([一二三四五六七八九十]+)[)）][ \s]*(.*)$'  # （一）xxx
    SUB_ITEM_PATTERN = r'^(\d+)[.、．][ \s]*(.*)$'  # 1. xxx 或 1、xxx
    
    def __init__(self):
        self.regulation = RegulationStructure()
    
    def parse(self, pages: List[Dict]) -> RegulationStructure:
        """
        解析法规结构
        
        Args:
            pages: PDF提取的页面数据
            
        Returns:
            法规结构化数据
        """
        # 合并所有文本
        full_text = "\n".join([p["text"] for p in pages])
        lines = full_text.split("\n")
        
        # 提取标题（通常在文档开头）
        self._extract_title(lines)
        
        # 解析章节结构
        self._parse_structure(lines)
        
        return self.regulation
    
    def _extract_title(self, lines: List[str]) -> None:
        """提取法规标题"""
        for line in lines[:20]:  # 通常标题在前20行
            line = line.strip()
            if line and len(line) > 5 and not line.startswith("第"):
                # 检查是否像法规标题
                if any(kw in line for kw in ["办法", "规定", "条例", "通知", "决定"]):
                    self.regulation.title = line
                    break
    
    def _parse_structure(self, lines: List[str]) -> None:
        """解析法规结构"""
        current_chapter = None
        current_article = None
        current_item = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检查章
            chapter_match = re.match(self.CHAPTER_PATTERN, line)
            if chapter_match:
                chapter_num = chapter_match.group(1)
                chapter_title = chapter_match.group(2).strip()
                current_chapter = {
                    "level": 1,
                    "number": f"第{chapter_num}章",
                    "title": chapter_title,
                    "content": "",
                    "children": []
                }
                self.regulation.chapters.append(current_chapter)
                current_article = None
                current_item = None
                continue
            
            # 检查条
            article_match = re.match(self.ARTICLE_PATTERN, line)
            if article_match:
                article_num = article_match.group(1)
                article_title = article_match.group(2).strip()
                
                current_article = {
                    "level": 2,
                    "number": f"第{article_num}条",
                    "title": article_title,
                    "content": "",
                    "children": []
                }
                
                if current_chapter:
                    current_chapter["children"].append(current_article)
                else:
                    # 没有章的情况，直接添加
                    self.regulation.chapters.append(current_article)
                
                current_item = None
                continue
            
            # 检查款
            item_match = re.match(self.ITEM_PATTERN, line)
            if item_match:
                item_num = item_match.group(1)
                item_title = item_match.group(2).strip()
                
                current_item = {
                    "level": 3,
                    "number": f"（{item_num}）",
                    "title": item_title,
                    "content": "",
                    "children": []
                }
                
                if current_article:
                    current_article["children"].append(current_item)
                
                continue
            
            # 检查目
            sub_item_match = re.match(self.SUB_ITEM_PATTERN, line)
            if sub_item_match:
                sub_num = sub_item_match.group(1)
                sub_title = sub_item_match.group(2).strip()
                
                sub_item = {
                    "level": 4,
                    "number": f"{sub_num}.",
                    "title": sub_title,
                    "content": "",
                    "children": []
                }
                
                if current_item:
                    current_item["children"].append(sub_item)
                elif current_article:
                    current_article["children"].append(sub_item)
                
                continue
            
            # 其他文本，追加到当前项
            if current_item:
                if current_item["content"]:
                    current_item["content"] += "\n" + line
                else:
                    current_item["content"] = line
            elif current_article:
                if current_article["content"]:
                    current_article["content"] += "\n" + line
                else:
                    current_article["content"] = line


class WordGenerator:
    """Word文档生成器"""
    
    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("请安装 python-docx: pip install python-docx")
        
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置标题样式
        styles = self.doc.styles
        
        # 标题1样式（章）
        if hasattr(styles, 'Heading 1'):
            heading1 = styles['Heading 1']
            heading1.font.size = Pt(16)
            heading1.font.bold = True
        
        # 标题2样式（条）
        if hasattr(styles, 'Heading 2'):
            heading2 = styles['Heading 2']
            heading2.font.size = Pt(14)
            heading2.font.bold = True
        
        # 标题3样式（款）
        if hasattr(styles, 'Heading 3'):
            heading3 = styles['Heading 3']
            heading3.font.size = Pt(12)
            heading3.font.bold = True
    
    def generate(self, regulation: RegulationStructure, output_path: str) -> None:
        """
        生成Word文档
        
        Args:
            regulation: 法规结构化数据
            output_path: 输出文件路径
        """
        # 添加标题
        if regulation.title:
            title = self.doc.add_heading(regulation.title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加章节内容
        for chapter in regulation.chapters:
            self._add_chapter(chapter)
        
        # 保存文档
        self.doc.save(output_path)
        print(f"已生成Word文档: {output_path}")
    
    def _add_chapter(self, chapter: Dict) -> None:
        """添加章节"""
        # 章标题
        if chapter["level"] == 1:
            # 一级标题（章）
            heading = self.doc.add_heading(
                f"{chapter['number']} {chapter['title']}", 
                level=1
            )
        else:
            # 可能有未分章的条
            self._add_article(chapter)
            return
        
        # 章内容
        if chapter.get("content"):
            self.doc.add_paragraph(chapter["content"])
        
        # 子项（条）
        for article in chapter.get("children", []):
            self._add_article(article)
    
    def _add_article(self, article: Dict) -> None:
        """添加条款"""
        # 条标题
        heading = self.doc.add_heading(
            f"{article['number']} {article['title']}", 
            level=2
        )
        
        # 条内容
        if article.get("content"):
            self.doc.add_paragraph(article["content"])
        
        # 子项（款）
        for item in article.get("children", []):
            self._add_item(item)
    
    def _add_item(self, item: Dict) -> None:
        """添加款项"""
        # 款标题
        para = self.doc.add_paragraph()
        run = para.add_run(f"{item['number']} {item['title']}")
        run.bold = True
        
        # 款内容
        if item.get("content"):
            self.doc.add_paragraph(item["content"])
        
        # 子项（目）
        for sub_item in item.get("children", []):
            self._add_sub_item(sub_item)
    
    def _add_sub_item(self, sub_item: Dict) -> None:
        """添加目"""
        para = self.doc.add_paragraph(style='List Bullet')
        para.add_run(f"{sub_item['number']} {sub_item['title']}")
        
        if sub_item.get("content"):
            self.doc.add_paragraph(sub_item["content"])


def pdf_to_word(pdf_path: str, output_path: Optional[str] = None, 
                 json_output: Optional[str] = None,
                 use_library: str = "auto") -> Tuple[str, str]:
    """
    将PDF法规文档转换为Word格式
    
    Args:
        pdf_path: PDF文件路径
        output_path: Word输出路径（可选，默认与PDF同名）
        json_output: JSON输出路径（可选）
        use_library: PDF解析库选择
        
    Returns:
        (word_path, json_path) 生成的文件路径
    """
    # 检查输入文件
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
    
    # 设置输出路径
    pdf_dir = os.path.dirname(pdf_path)
    pdf_name = Path(pdf_path).stem
    
    if output_path is None:
        output_path = os.path.join(pdf_dir, f"{pdf_name}.docx")
    
    if json_output is None:
        json_output = os.path.join(pdf_dir, f"{pdf_name}.json")
    
    print(f"正在解析PDF: {pdf_path}")
    
    # 1. 提取PDF文本
    parser = PDFParser(use_library)
    pages = parser.extract_text(pdf_path)
    print(f"  - 提取了 {len(pages)} 页")
    
    # 2. 解析法规结构
    reg_parser = RegulationParser()
    regulation = reg_parser.parse(pages)
    print(f"  - 解析出 {len(regulation.chapters)} 个章节")
    
    # 3. 生成Word文档
    word_gen = WordGenerator()
    word_gen.generate(regulation, output_path)
    
    # 4. 导出JSON（可选）
    regulation.to_json(json_output)
    
    return output_path, json_output


def batch_convert(input_dir: str, output_dir: Optional[str] = None) -> List[Dict]:
    """
    批量转换PDF文件
    
    Args:
        input_dir: 输入目录
        output_dir: 输出目录（可选，默认与输入目录相同）
        
    Returns:
        转换结果列表
    """
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    results = []
    pdf_files = list(Path(input_dir).glob("*.pdf"))
    
    print(f"找到 {len(pdf_files)} 个PDF文件")
    
    for pdf_file in pdf_files:
        try:
            pdf_path = str(pdf_file)
            
            if output_dir:
                word_path = os.path.join(output_dir, f"{pdf_file.stem}.docx")
                json_path = os.path.join(output_dir, f"{pdf_file.stem}.json")
            else:
                word_path = None
                json_path = None
            
            word_path, json_path = pdf_to_word(pdf_path, word_path, json_path)
            
            results.append({
                "pdf": pdf_path,
                "word": word_path,
                "json": json_path,
                "status": "success"
            })
            
        except Exception as e:
            print(f"转换失败 {pdf_file}: {e}")
            results.append({
                "pdf": str(pdf_file),
                "status": "failed",
                "error": str(e)
            })
    
    return results


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(description="PDF法规文档转Word工具")
    parser.add_argument("input", help="PDF文件路径或目录")
    parser.add_argument("-o", "--output", help="输出路径或目录")
    parser.add_argument("-l", "--library", default="auto", 
                        choices=["auto", "pymupdf", "pdfplumber"],
                        help="PDF解析库选择")
    parser.add_argument("--no-json", action="store_true", help="不生成JSON文件")
    
    args = parser.parse_args()
    
    input_path = args.input
    
    if os.path.isdir(input_path):
        # 批量转换
        results = batch_convert(input_path, args.output)
        
        success = sum(1 for r in results if r["status"] == "success")
        failed = sum(1 for r in results if r["status"] == "failed")
        
        print(f"\n转换完成: 成功 {success}, 失败 {failed}")
        
    elif os.path.isfile(input_path):
        # 单文件转换
        json_output = None if args.no_json else args.output and args.output.replace(".docx", ".json")
        
        word_path, json_path = pdf_to_word(
            input_path, 
            args.output,
            json_output,
            args.library
        )
        
        print(f"\n转换完成:")
        print(f"  Word: {word_path}")
        if json_path:
            print(f"  JSON: {json_path}")
    else:
        print(f"错误: 路径不存在 {input_path}")


if __name__ == "__main__":
    main()
